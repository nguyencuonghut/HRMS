"""Tests cho 4.2 — sinh hợp đồng từ mẫu DOCX."""
from __future__ import annotations

import asyncio
from datetime import date
from decimal import Decimal
from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import delete, select, text
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from app.core.config import settings
from app.models.employee import Employee
from app.services.contract_generate_service import (
    fmt_vn_date,
    number_to_words_vn,
    render_contract_docx,
)

BASE_EMP      = "/api/v1/employees"
BASE_TEMPLATE = "/api/v1/contract-templates"

_ADMIN_EMAIL    = "admin@hrms.local"
_ADMIN_PASSWORD = "Hrms@2026"
_LM_EMAIL       = "linemanager@hrms.local"

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_PREFIX    = "TESTGEN"

# Seed category IDs (labor_indefinite = 1, labor_definite = 2, appendix_salary = 3)
CAT_LABOR_INDEFINITE = 1
CAT_APPENDIX_SALARY  = 3


# ── Auth helpers ───────────────────────────────────────────────────────────────

def _login(client: TestClient, email: str = _ADMIN_EMAIL, pw: str = _ADMIN_PASSWORD) -> dict:
    token = client.post("/api/v1/auth/login", json={"email": email, "password": pw}).json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def _admin(client: TestClient) -> dict: return _login(client)
def _lm(client: TestClient)   -> dict: return _login(client, _LM_EMAIL)


# ── Cleanup ────────────────────────────────────────────────────────────────────

def _make_session():
    engine = create_async_engine(settings.DATABASE_URL, connect_args={"ssl": False})
    return async_sessionmaker(engine, expire_on_commit=False)


async def _cleanup():
    async with _make_session()() as s:
        employee_ids = [e.id for e in (await s.execute(select(Employee))).scalars().all() if e.id_number.startswith(_PREFIX)]
        await s.execute(text(
            f"DELETE FROM employee_contracts WHERE contract_number LIKE '{_PREFIX}%'"
        ))
        if employee_ids:
            await s.execute(delete(Employee).where(Employee.id.in_(employee_ids)))
        await s.execute(text(
            "DELETE FROM contract_template_placeholders WHERE template_id IN "
            f"(SELECT id FROM contract_templates WHERE code LIKE '{_PREFIX.lower()}%')"
        ))
        await s.execute(text(
            f"DELETE FROM contract_templates WHERE code LIKE '{_PREFIX.lower()}%'"
        ))
        await s.commit()


@pytest.fixture(autouse=True)
async def cleanup():
    yield
    await _cleanup()


# ── DOCX builder helper ────────────────────────────────────────────────────────

def _make_docx_bytes(*placeholders: str) -> bytes:
    """Tạo DOCX tối thiểu với các placeholder {{...}}."""
    doc = Document()
    for ph in placeholders:
        doc.add_paragraph(ph)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_split_run_docx(left: str, right: str) -> bytes:
    """DOCX với placeholder bị tách ra làm 2 run."""
    doc = Document()
    para = doc.add_paragraph()
    para.add_run(left)
    para.add_run(right)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Setup helpers ──────────────────────────────────────────────────────────────

def _create_employee(client: TestClient, headers: dict, suffix: str) -> int:
    resp = client.post(BASE_EMP, json={
        "employee_code_sequence_id": 1,
        "full_name": f"Test Gen {suffix}",
        "last_name": "Test",
        "first_name": f"Gen {suffix}",
        "date_of_birth": "1990-06-15",
        "gender": "male",
        "nationality_id": 1,
        "id_number": f"{_PREFIX}{suffix}",
        "id_issued_on": "2020-01-01",
        "id_issued_by": "Cục CA",
        "status": "official",
        "start_date": "2020-01-01",
    }, headers=headers)
    assert resp.status_code == 201, resp.text
    return resp.json()["id"]


def _create_contract(client: TestClient, headers: dict, emp_id: int, suffix: str,
                     category_id: int = CAT_LABOR_INDEFINITE, **kwargs) -> dict:
    today = date.today().isoformat()
    payload = {
        "contract_category_id": category_id,
        "contract_number": f"{_PREFIX}-HĐ-{suffix}",
        "signed_date": today,
        "effective_from": today,
        "effective_to": None,
        "insurance_salary": "8000000",
    }
    payload.update(kwargs)
    resp = client.post(f"{BASE_EMP}/{emp_id}/contracts", json=payload, headers=headers)
    assert resp.status_code == 201, resp.text
    return resp.json()


def _create_template(client: TestClient, headers: dict, suffix: str,
                     document_kind: str = "labor_contract") -> int:
    """Tạo ContractTemplate record (không có file)."""
    resp = client.post(BASE_TEMPLATE, json={
        "code": f"{_PREFIX.lower()}_tpl_{suffix}",
        "name": f"Mẫu test {suffix}",
        "contract_category_id": CAT_LABOR_INDEFINITE,
        "document_kind": document_kind,
        "template_engine": "docx_placeholders",
        "file_name": f"test_{suffix}.docx",
        "mime_type": _DOCX_MIME,
        "version_no": 1,
    }, headers=headers)
    assert resp.status_code == 201, resp.text
    return resp.json()["id"]


def _upload_docx(client: TestClient, headers: dict, template_id: int,
                 docx_bytes: bytes, filename: str = "test.docx") -> dict:
    resp = client.post(
        f"{BASE_TEMPLATE}/{template_id}/upload",
        files={"file": (filename, docx_bytes, _DOCX_MIME)},
        headers=headers,
    )
    assert resp.status_code == 200, resp.text
    return resp.json()


# ── Unit tests — formatters & render ──────────────────────────────────────────

def test_fmt_vn_date():
    assert fmt_vn_date(date(2024, 1, 15)) == "15 tháng 01 năm 2024"
    assert fmt_vn_date(date(2024, 12, 31)) == "31 tháng 12 năm 2024"
    assert fmt_vn_date(None) == ""


def test_number_to_words_vn():
    assert number_to_words_vn(Decimal("0")) == "Không đồng"
    assert number_to_words_vn(Decimal("8000000")) == "Tám triệu đồng"
    assert number_to_words_vn(Decimal("15000000")) == "Mười lăm triệu đồng"
    assert number_to_words_vn(Decimal("5500000")) == "Năm triệu năm trăm nghìn đồng"
    assert number_to_words_vn(Decimal("1200000")) == "Một triệu hai trăm nghìn đồng"
    assert number_to_words_vn(Decimal("10000000")) == "Mười triệu đồng"
    assert number_to_words_vn(None) == ""


def test_render_docx_replaces_tokens():
    tmpl = _make_docx_bytes(
        "Họ và tên: {{employee_full_name}}",
        "Số HĐ: {{contract_number}}",
        "Không có token ở đây",
    )
    ctx = {"employee_full_name": "Nguyễn Văn A", "contract_number": "HĐ-2024-001"}
    result = render_contract_docx(tmpl, ctx)
    doc = Document(BytesIO(result))
    texts = [p.text for p in doc.paragraphs]
    assert "Họ và tên: Nguyễn Văn A" in texts
    assert "Số HĐ: HĐ-2024-001" in texts
    assert "Không có token ở đây" in texts


def test_render_docx_split_run_placeholder():
    """Placeholder bị tách qua 2 run phải được thay đúng."""
    tmpl = _make_split_run_docx("Ngày sinh: {{employee_", "birthday}}")
    ctx = {"employee_birthday": "15 tháng 06 năm 1990"}
    result = render_contract_docx(tmpl, ctx)
    doc = Document(BytesIO(result))
    texts = [p.text for p in doc.paragraphs]
    assert any("Ngày sinh: 15 tháng 06 năm 1990" in t for t in texts), texts


def test_render_docx_unknown_token_replaced_with_empty():
    tmpl = _make_docx_bytes("Giá trị: {{unknown_token}}")
    result = render_contract_docx(tmpl, {})
    doc = Document(BytesIO(result))
    assert any("Giá trị: " == p.text for p in doc.paragraphs)


# ── Upload template tests ──────────────────────────────────────────────────────

def test_upload_template_docx_success(client: TestClient):
    headers = _admin(client)
    tid = _create_template(client, headers, "UP01")
    docx = _make_docx_bytes("{{employee_full_name}}")
    data = _upload_docx(client, headers, tid, docx)
    assert data["storage_path"] is not None
    assert data["storage_path"].startswith(f"templates/{tid}/")
    assert data["file_size"] > 0
    assert data["file_checksum"] is not None


def test_upload_template_invalid_type_400(client: TestClient):
    headers = _admin(client)
    tid = _create_template(client, headers, "UP02")
    resp = client.post(
        f"{BASE_TEMPLATE}/{tid}/upload",
        files={"file": ("report.pdf", b"%PDF-1.4 content", "application/pdf")},
        headers=headers,
    )
    assert resp.status_code == 400
    assert "docx" in resp.json()["detail"].lower()


def test_upload_template_too_large_400(client: TestClient):
    headers = _admin(client)
    tid = _create_template(client, headers, "UP03")
    big = b"x" * (5 * 1024 * 1024 + 1)
    resp = client.post(
        f"{BASE_TEMPLATE}/{tid}/upload",
        files={"file": ("big.docx", big, _DOCX_MIME)},
        headers=headers,
    )
    assert resp.status_code == 400
    assert "5 MB" in resp.json()["detail"]


def test_download_template_success(client: TestClient):
    headers = _admin(client)
    tid = _create_template(client, headers, "DL01")
    docx = _make_docx_bytes("{{contract_number}}")
    _upload_docx(client, headers, tid, docx)

    resp = client.get(f"{BASE_TEMPLATE}/{tid}/download", headers=headers)
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(_DOCX_MIME)
    assert len(resp.content) > 0


def test_download_template_no_file_404(client: TestClient):
    headers = _admin(client)
    tid = _create_template(client, headers, "DL02")  # no upload
    resp = client.get(f"{BASE_TEMPLATE}/{tid}/download", headers=headers)
    assert resp.status_code == 404


# ── Generate endpoint tests ────────────────────────────────────────────────────

def test_generate_contract_success(client: TestClient):
    headers = _admin(client)
    emp_id = _create_employee(client, headers, "G01")
    con    = _create_contract(client, headers, emp_id, "G01")
    tid    = _create_template(client, headers, "G01")
    docx   = _make_docx_bytes(
        "Tên: {{employee_full_name}}",
        "Số HĐ: {{contract_number}}",
        "Lương: {{insurance_salary}}",
    )
    _upload_docx(client, headers, tid, docx)

    resp = client.post(
        f"{BASE_EMP}/{emp_id}/contracts/{con['id']}/generate",
        json={"template_id": tid},
        headers=headers,
    )
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith(_DOCX_MIME)
    assert len(resp.content) > 0

    # Nội dung DOCX phải chứa tên nhân viên
    doc = Document(BytesIO(resp.content))
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "Test Gen G01" in texts
    assert f"{_PREFIX}-HĐ-G01" in texts


def test_generate_contract_wrong_employee_404(client: TestClient):
    headers = _admin(client)
    emp1 = _create_employee(client, headers, "G02A")
    emp2 = _create_employee(client, headers, "G02B")
    con  = _create_contract(client, headers, emp1, "G02")
    tid  = _create_template(client, headers, "G02")
    _upload_docx(client, headers, tid, _make_docx_bytes("{{contract_number}}"))

    # Dùng emp2 nhưng contract thuộc emp1
    resp = client.post(
        f"{BASE_EMP}/{emp2}/contracts/{con['id']}/generate",
        json={"template_id": tid},
        headers=headers,
    )
    assert resp.status_code == 404


def test_generate_contract_terminated_400(client: TestClient):
    headers = _admin(client)
    emp_id = _create_employee(client, headers, "G03")
    con    = _create_contract(client, headers, emp_id, "G03")
    tid    = _create_template(client, headers, "G03")
    _upload_docx(client, headers, tid, _make_docx_bytes("{{employee_full_name}}"))

    # Chấm dứt hợp đồng
    client.delete(f"{BASE_EMP}/{emp_id}/contracts/{con['id']}", headers=headers)

    resp = client.post(
        f"{BASE_EMP}/{emp_id}/contracts/{con['id']}/generate",
        json={"template_id": tid},
        headers=headers,
    )
    assert resp.status_code == 400
    assert "chấm dứt" in resp.json()["detail"]


def test_generate_contract_no_template_file_422(client: TestClient):
    headers = _admin(client)
    emp_id = _create_employee(client, headers, "G04")
    con    = _create_contract(client, headers, emp_id, "G04")
    tid    = _create_template(client, headers, "G04")  # no upload

    resp = client.post(
        f"{BASE_EMP}/{emp_id}/contracts/{con['id']}/generate",
        json={"template_id": tid},
        headers=headers,
    )
    assert resp.status_code == 422


def test_generate_contract_kind_mismatch_400(client: TestClient):
    headers = _admin(client)
    emp_id = _create_employee(client, headers, "G05")
    # HĐ gốc (labor_contract)
    con    = _create_contract(client, headers, emp_id, "G05")
    # Mẫu dành cho phụ lục (contract_appendix)
    tid    = _create_template(client, headers, "G05", document_kind="contract_appendix")
    _upload_docx(client, headers, tid, _make_docx_bytes("{{contract_number}}"))

    resp = client.post(
        f"{BASE_EMP}/{emp_id}/contracts/{con['id']}/generate",
        json={"template_id": tid},
        headers=headers,
    )
    assert resp.status_code == 400
    assert "không khớp" in resp.json()["detail"]


def test_generate_contract_unauth_401(client: TestClient):
    resp = client.post(
        f"{BASE_EMP}/1/contracts/1/generate",
        json={"template_id": 1},
    )
    assert resp.status_code == 401


def test_generate_contract_no_perm_403(client: TestClient):
    headers = _lm(client)
    resp = client.post(
        f"{BASE_EMP}/1/contracts/1/generate",
        json={"template_id": 1},
        headers=headers,
    )
    assert resp.status_code == 403
