"""Dịch vụ sinh hợp đồng từ mẫu DOCX — 4.2."""
from __future__ import annotations

import re
from datetime import date
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import Optional, Union

from docx import Document
from fastapi import HTTPException, status
from sqlalchemy.ext.asyncio import AsyncSession
from sqlmodel import select

from app.models.catalog import ContractCategory, ContractTemplate
from app.models.employee import Employee, EmployeeAddress
from app.models.employee_contract import EmployeeContract
from app.models.employee_job import EmployeeJobRecord
from app.models.org import Department, JobTitle

_DOUBLE_PATTERN = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")
_DOLLAR_PATTERN = re.compile(r"\$\{\s*([^}]+?)\s*\}")

# ── Formatters ────────────────────────────────────────────────────────────────

def fmt_vn_date(d: Optional[date]) -> str:
    """date(2024, 1, 15) → '15 tháng 01 năm 2024'"""
    if d is None:
        return ""
    return f"{d.day:02d} tháng {d.month:02d} năm {d.year}"


def fmt_currency_vnd(v: Optional[Decimal]) -> str:
    """Decimal('8000000') → '8.000.000 đồng'"""
    if v is None:
        return ""
    return f"{int(v):,}".replace(",", ".") + " đồng"


def gender_label(g: Optional[str]) -> str:
    """'male' → 'Nam', 'female' → 'Nữ'"""
    if g == "male":
        return "Nam"
    if g == "female":
        return "Nữ"
    return g or ""


_ONES = ["không", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]


def _three_digits(n: int, has_higher: bool = False) -> str:
    """Đọc 0-999 thành chữ tiếng Việt. has_higher=True khi nhóm cao hơn != 0."""
    if n == 0:
        return "không trăm" if has_higher else ""
    h, r = divmod(n, 100)
    t, u = divmod(r, 10)
    parts: list[str] = []
    if h:
        parts.append(f"{_ONES[h]} trăm")
    elif has_higher:
        parts.append("không trăm")
    if t == 0:
        if u:
            parts.append(f"linh {_ONES[u]}" if (h or has_higher) else _ONES[u])
    elif t == 1:
        s = "mười"
        if u == 5:
            s += " lăm"
        elif u:
            s += f" {_ONES[u]}"
        parts.append(s)
    else:
        s = f"{_ONES[t]} mươi"
        if u == 1:
            s += " mốt"
        elif u == 4:
            s += " tư"
        elif u == 5:
            s += " lăm"
        elif u:
            s += f" {_ONES[u]}"
        parts.append(s)
    return " ".join(parts)


def number_to_words_vn(v: Optional[Decimal]) -> str:
    """Decimal('8000000') → 'Tám triệu đồng'"""
    if v is None:
        return ""
    n = int(v)
    if n == 0:
        return "Không đồng"
    negative = n < 0
    n = abs(n)
    ty  = n // 1_000_000_000
    tr  = (n % 1_000_000_000) // 1_000_000
    ng  = (n % 1_000_000) // 1_000
    dv  = n % 1_000
    parts: list[str] = []
    if ty:
        parts.append(f"{_three_digits(ty)} tỷ")
    if tr:
        parts.append(f"{_three_digits(tr, bool(ty))} triệu")
    if ng:
        parts.append(f"{_three_digits(ng, bool(ty or tr))} nghìn")
    if dv:
        parts.append(_three_digits(dv, bool(ty or tr or ng)))
    result = " ".join(parts)
    if negative:
        result = "âm " + result
    return result[0].upper() + result[1:] + " đồng"


# ── DOCX rendering ────────────────────────────────────────────────────────────

def _iter_paragraphs(doc: Document):
    """Yield tất cả paragraph trong body + table + header/footer."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for hdr_ftr in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            try:
                if not hdr_ftr.is_linked_to_previous:
                    yield from hdr_ftr.paragraphs
            except Exception:
                pass


def _replace_paragraph(paragraph, context: dict[str, str]) -> None:
    """Gộp toàn bộ runs → thay placeholder ({{}} và ${}) → ghi vào run[0], xóa run còn lại."""
    if not paragraph.runs:
        return
    full_text = "".join(r.text for r in paragraph.runs)
    if "{{" not in full_text and "${" not in full_text:
        return
    replaced = _DOUBLE_PATTERN.sub(lambda m: context.get(m.group(1).strip(), ""), full_text)
    replaced = _DOLLAR_PATTERN.sub(lambda m: context.get(m.group(1).strip(), ""), replaced)
    paragraph.runs[0].text = replaced
    for run in paragraph.runs[1:]:
        run.text = ""


def render_contract_docx(
    template_source: Union[Path, bytes],
    context: dict[str, str],
) -> bytes:
    """Render DOCX: thay {{token}} bằng context values, trả về bytes."""
    if isinstance(template_source, bytes):
        doc = Document(BytesIO(template_source))
    else:
        doc = Document(str(template_source))
    for para in _iter_paragraphs(doc):
        _replace_paragraph(para, context)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Data resolver ─────────────────────────────────────────────────────────────

async def build_contract_context(
    session: AsyncSession,
    employee_id: int,
    contract_id: int,
) -> dict[str, str]:
    """Load Employee + EmployeeContract + job record → dict token: str_value."""
    emp = await session.get(Employee, employee_id)
    if not emp:
        raise HTTPException(status.HTTP_404_NOT_FOUND, detail="Nhân viên không tồn tại")

    contract = await session.get(EmployeeContract, contract_id)
    if not contract or contract.employee_id != employee_id:
        raise HTTPException(status.HTTP_404_NOT_FOUND, detail="Hợp đồng không tồn tại")

    category = await session.get(ContractCategory, contract.contract_category_id)

    # Current job record
    job_result = await session.execute(
        select(EmployeeJobRecord).where(
            EmployeeJobRecord.employee_id == employee_id,
            EmployeeJobRecord.is_current == True,  # noqa: E712
        )
    )
    job = job_result.scalars().first()
    dept_name = ""
    job_title_name = ""
    if job:
        dept = await session.get(Department, job.department_id)
        dept_name = dept.name if dept else ""
        if job.job_title_id:
            jt = await session.get(JobTitle, job.job_title_id)
            job_title_name = jt.name if jt else ""

    # Addresses
    addr_result = await session.execute(
        select(EmployeeAddress).where(EmployeeAddress.employee_id == employee_id)
    )
    addresses = {a.address_type: a for a in addr_result.scalars().all()}
    perm = addresses.get("permanent")
    contact = addresses.get("contact")

    sign_date = contract.signed_date

    return {
        "contract_number":        contract.contract_number or "",
        "contract_start_date":    fmt_vn_date(contract.effective_from),
        "contract_end_date":      fmt_vn_date(contract.effective_to),
        "department_name":        dept_name,
        "employee_full_name":     emp.full_name or "",
        "employee_birthday":      fmt_vn_date(emp.date_of_birth),
        "employee_gender":        gender_label(emp.gender),
        "employee_cccd":          emp.id_number or "",
        "employee_cccd_issued_on": fmt_vn_date(emp.id_issued_on),
        "employee_cccd_issued_by": emp.id_issued_by or "",
        "employee_address":       perm.full_address_text if perm else "",
        "employee_temp_address":  contact.full_address_text if contact else "",
        "employee_phone":         emp.phone_number or "",
        "employee_personal_email": emp.personal_email or "",
        "position_title":         job_title_name,
        "insurance_salary":       fmt_currency_vnd(contract.insurance_salary),
        "insurance_salary_words": number_to_words_vn(contract.insurance_salary),
        # Legacy tokens
        "Ngày":             str(sign_date.day)   if sign_date else "",
        "Tháng":            str(sign_date.month) if sign_date else "",
        "Năm":              str(sign_date.year)  if sign_date else "",
        "SĐT":              emp.phone_number or "",
        "Loại_HĐLĐ__":     category.name if category else "",
        "Thời_hạn_trả_lương": "",
    }


# ── Orchestrator ──────────────────────────────────────────────────────────────

async def generate_contract_document(
    session: AsyncSession,
    employee_id: int,
    contract_id: int,
    template_id: int,
) -> tuple[bytes, str]:
    """Load template → build context → render DOCX → (bytes, filename)."""
    from app.core.storage import get_object_bytes

    template = await session.get(ContractTemplate, template_id)
    if not template or not template.is_active:
        raise HTTPException(
            status.HTTP_404_NOT_FOUND,
            detail="Mẫu hợp đồng không tồn tại hoặc đã vô hiệu hóa",
        )
    if not template.storage_path:
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Mẫu chưa có file DOCX đính kèm",
        )

    contract = await session.get(EmployeeContract, contract_id)
    if not contract or contract.employee_id != employee_id:
        raise HTTPException(status.HTTP_404_NOT_FOUND, detail="Hợp đồng không tồn tại")
    if contract.status == "terminated":
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            detail="Không thể sinh hợp đồng đã chấm dứt",
        )
    if template.document_kind != contract.document_kind:
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            detail=f"Loại mẫu '{template.document_kind}' không khớp với loại hợp đồng '{contract.document_kind}'",
        )

    # Thử MinIO trước; fallback sang đường dẫn local (seed data)
    try:
        docx_source: Union[Path, bytes] = get_object_bytes(template.storage_path)
    except Exception:
        from app.services.contract_template_docx import resolve_template_storage_path
        local_path = resolve_template_storage_path(template.storage_path)
        if not local_path.exists():
            raise HTTPException(
                status.HTTP_404_NOT_FOUND,
                detail=f"Không tìm thấy file mẫu tại '{template.storage_path}'",
            )
        docx_source = local_path

    context = await build_contract_context(session, employee_id, contract_id)
    rendered = render_contract_docx(docx_source, context)
    filename = f"HD_{contract.contract_number}.docx"
    return rendered, filename
